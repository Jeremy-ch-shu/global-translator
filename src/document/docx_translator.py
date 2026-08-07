from docx import Document
from docx.shared import Pt, RGBColor
from loguru import logger

def translate_docx(file_path: str, engine, target_lang: str = 'zh-CN'):
    """翻译Word文档并保留格式"""
    
    # 读取原文档
    doc = Document(file_path)
    
    # 遍历所有段落
    for para in doc.paragraphs:
        if para.text.strip():
            # 提取每个run的格式信息[reference:19]
            runs_info = []
            for run in para.runs:
                runs_info.append({
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'color': run.font.color.rgb if run.font.color else None
                })
            
            # 翻译完整段落
            full_text = para.text
            translated = engine.translate(full_text, 'auto', target_lang)
            
            # 清空段落并重建
            para.clear()
            # 应用第一个run的格式到整个段落（简化处理）
            if runs_info:
                new_run = para.add_run(translated)
                info = runs_info[0]
                new_run.bold = info['bold']
                new_run.italic = info['italic']
                new_run.underline = info['underline']
                if info['font_name']:
                    new_run.font.name = info['font_name']
                if info['font_size']:
                    new_run.font.size = info['font_size']
                if info['color']:
                    new_run.font.color.rgb = info['color']
            else:
                para.add_run(translated)
    
    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        translated = engine.translate(para.text, 'auto', target_lang)
                        para.clear()
                        para.add_run(translated)
    
    # 保存新文档
    output_path = file_path.replace('.docx', '_translated.docx')
    doc.save(output_path)
    logger.info(f"文档翻译完成: {output_path}")
    return output_path